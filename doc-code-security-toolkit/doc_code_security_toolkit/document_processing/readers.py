"""Document readers for various formats.

This module provides readers for extracting text, tables, and metadata
from different document formats including PDF, Word, Excel, and more.
"""

import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import pandas as pd


@dataclass
class TableData:
    """Represents extracted table data."""

    table_id: str
    page_number: Optional[int] = None
    headers: List[str] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)
    data_frame: Optional[pd.DataFrame] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "table_id": self.table_id,
            "page_number": self.page_number,
            "headers": self.headers,
            "rows": self.rows,
            "row_count": len(self.rows),
            "column_count": len(self.headers) if self.headers else (len(self.rows[0]) if self.rows else 0),
            "metadata": self.metadata,
        }

    def to_dataframe(self) -> pd.DataFrame:
        """Convert to pandas DataFrame."""
        if self.data_frame is not None:
            return self.data_frame
        if self.headers:
            return pd.DataFrame(self.rows, columns=self.headers)
        return pd.DataFrame(self.rows)


@dataclass
class DocumentContent:
    """Represents extracted document content."""

    file_path: str
    file_name: str
    file_size: int
    document_type: str
    page_count: int = 0
    text: str = ""
    text_by_page: Dict[int, str] = field(default_factory=dict)
    tables: List[TableData] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    extracted_at: datetime = field(default_factory=datetime.now)
    images: List[Dict[str, Any]] = field(default_factory=list)
    links: List[Dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "file_path": self.file_path,
            "file_name": self.file_name,
            "file_size": self.file_size,
            "document_type": self.document_type,
            "page_count": self.page_count,
            "text_length": len(self.text),
            "text_preview": self.text[:500] + "..." if len(self.text) > 500 else self.text,
            "table_count": len(self.tables),
            "tables": [t.to_dict() for t in self.tables],
            "metadata": self.metadata,
            "extracted_at": self.extracted_at.isoformat(),
            "image_count": len(self.images),
            "link_count": len(self.links),
        }

    def save_text(self, output_path: str) -> Path:
        """Save extracted text to file."""
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(self.text)
        return path

    def save_tables(
        self,
        output_dir: str,
        format: str = "csv",
    ) -> List[Path]:
        """Save extracted tables to files.

        Args:
            output_dir: Output directory
            format: Output format (csv, excel, json)

        Returns:
            List of saved file paths
        """
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        saved_paths = []

        for idx, table in enumerate(self.tables):
            df = table.to_dataframe()
            base_name = f"table_{idx + 1}_{table.table_id}"

            if format == "csv":
                file_path = output_path / f"{base_name}.csv"
                df.to_csv(file_path, index=False, encoding="utf-8")
            elif format == "excel":
                file_path = output_path / f"{base_name}.xlsx"
                df.to_excel(file_path, index=False)
            elif format == "json":
                file_path = output_path / f"{base_name}.json"
                df.to_json(file_path, orient="records", indent=2)
            else:
                raise ValueError(f"Unsupported format: {format}")

            saved_paths.append(file_path)

        return saved_paths


class DocumentReader(ABC):
    """Abstract base class for document readers."""

    supported_formats: List[str] = []

    @abstractmethod
    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read and extract content from document.

        Args:
            file_path: Path to the document file
            extract_tables: Whether to extract tables
            extract_images: Whether to extract images
            extract_links: Whether to extract links

        Returns:
            DocumentContent object with extracted content
        """
        pass

    @classmethod
    def supports_format(cls, file_extension: str) -> bool:
        """Check if this reader supports the given file extension."""
        ext = file_extension.lower().lstrip(".")
        return ext in [f.lower().lstrip(".") for f in cls.supported_formats]


class PDFReader(DocumentReader):
    """Reader for PDF documents using pdfplumber."""

    supported_formats = ["pdf"]

    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read PDF document."""
        import pdfplumber

        path = Path(file_path)
        file_stat = path.stat()

        content = DocumentContent(
            file_path=str(path),
            file_name=path.name,
            file_size=file_stat.st_size,
            document_type="pdf",
        )

        try:
            with pdfplumber.open(path) as pdf:
                content.page_count = len(pdf.pages)

                all_text = []
                tables = []

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    content.text_by_page[page_num] = page_text
                    all_text.append(page_text)

                    if extract_tables:
                        page_tables = page.extract_tables()
                        for table_idx, table_data in enumerate(page_tables):
                            if table_data:
                                table = TableData(
                                    table_id=f"page{page_num}_table{table_idx + 1}",
                                    page_number=page_num,
                                )
                                if len(table_data) > 1:
                                    table.headers = table_data[0] if table_data else []
                                    table.rows = table_data[1:] if len(table_data) > 1 else []
                                else:
                                    table.rows = table_data

                                try:
                                    table.data_frame = pd.DataFrame(
                                        table.rows,
                                        columns=table.headers if table.headers else None,
                                    )
                                except Exception:
                                    pass

                                tables.append(table)

                content.text = "\n\n".join(all_text)
                content.tables = tables

                if pdf.metadata:
                    content.metadata = {
                        k: str(v) if v is not None else ""
                        for k, v in pdf.metadata.items()
                    }

        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF reading. "
                "Install it with: pip install pdfplumber"
            )

        return content


class WordReader(DocumentReader):
    """Reader for Word documents (.docx)."""

    supported_formats = ["docx", "docm"]

    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read Word document."""
        from docx import Document
        from docx.document import Document as DocxDocument
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph

        path = Path(file_path)
        file_stat = path.stat()

        content = DocumentContent(
            file_path=str(path),
            file_name=path.name,
            file_size=file_stat.st_size,
            document_type="word",
            page_count=1,
        )

        try:
            doc: DocxDocument = Document(path)

            all_text = []
            tables = []

            for para in doc.paragraphs:
                all_text.append(para.text)

            if extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = TableData(
                        table_id=f"table{table_idx + 1}",
                        page_number=1,
                    )

                    rows_data = []
                    for row_idx, row in enumerate(table.rows):
                        row_text = [cell.text for cell in row.cells]
                        rows_data.append(row_text)

                    if len(rows_data) > 1:
                        table_data.headers = rows_data[0]
                        table_data.rows = rows_data[1:]
                    else:
                        table_data.rows = rows_data

                    try:
                        table_data.data_frame = pd.DataFrame(
                            table_data.rows,
                            columns=table_data.headers if table_data.headers else None,
                        )
                    except Exception:
                        pass

                    tables.append(table_data)

            content.text = "\n\n".join(all_text)
            content.tables = tables

            core_props = doc.core_properties
            content.metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
            }

        except ImportError:
            raise ImportError(
                "python-docx is required for Word document reading. "
                "Install it with: pip install python-docx"
            )

        return content


class ExcelReader(DocumentReader):
    """Reader for Excel spreadsheets (.xlsx, .xls)."""

    supported_formats = ["xlsx", "xls", "xlsm"]

    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read Excel spreadsheet."""
        path = Path(file_path)
        file_stat = path.stat()

        content = DocumentContent(
            file_path=str(path),
            file_name=path.name,
            file_size=file_stat.st_size,
            document_type="excel",
            page_count=1,
        )

        try:
            xls = pd.ExcelFile(path)

            all_text = []
            tables = []

            for sheet_idx, sheet_name in enumerate(xls.sheet_names):
                df = pd.read_excel(xls, sheet_name=sheet_name)

                all_text.append(f"--- Sheet: {sheet_name} ---")
                all_text.append(df.to_string())

                if extract_tables:
                    table = TableData(
                        table_id=f"sheet_{sheet_name.replace(' ', '_')}",
                        page_number=sheet_idx + 1,
                    )
                    table.headers = list(df.columns)
                    table.rows = df.values.tolist()
                    table.data_frame = df.copy()
                    table.metadata = {
                        "sheet_name": sheet_name,
                        "dtypes": {str(k): str(v) for k, v in df.dtypes.items()},
                    }
                    tables.append(table)

            content.text = "\n\n".join(all_text)
            content.tables = tables
            content.metadata = {
                "sheet_count": len(xls.sheet_names),
                "sheet_names": xls.sheet_names,
            }

        except ImportError:
            raise ImportError(
                "openpyxl and pandas are required for Excel reading. "
                "Install them with: pip install openpyxl pandas"
            )

        return content


class TextReader(DocumentReader):
    """Reader for plain text files and Markdown."""

    supported_formats = ["txt", "md", "rst", "csv", "json"]

    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read text file."""
        path = Path(file_path)
        file_stat = path.stat()
        ext = path.suffix.lower().lstrip(".")

        content = DocumentContent(
            file_path=str(path),
            file_name=path.name,
            file_size=file_stat.st_size,
            document_type=ext,
            page_count=1,
        )

        tables = []

        if ext == "csv" and extract_tables:
            try:
                df = pd.read_csv(path)
                table = TableData(
                    table_id="csv_table",
                    page_number=1,
                )
                table.headers = list(df.columns)
                table.rows = df.values.tolist()
                table.data_frame = df.copy()
                tables.append(table)

                content.text = df.to_string()
            except Exception:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    content.text = f.read()
        elif ext == "json":
            import json

            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                if isinstance(data, list) and all(isinstance(item, dict) for item in data):
                    try:
                        df = pd.DataFrame(data)
                        table = TableData(
                            table_id="json_table",
                            page_number=1,
                        )
                        table.headers = list(df.columns)
                        table.rows = df.values.tolist()
                        table.data_frame = df.copy()
                        tables.append(table)
                        content.text = df.to_string()
                    except Exception:
                        content.text = json.dumps(data, indent=2)
                else:
                    content.text = json.dumps(data, indent=2)
            except Exception:
                with open(path, "r", encoding="utf-8", errors="replace") as f:
                    content.text = f.read()
        else:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                content.text = f.read()

            if ext == "md" and extract_tables:
                tables = self._extract_markdown_tables(content.text)

        content.tables = tables
        content.metadata = {
            "encoding": "utf-8",
            "line_count": content.text.count("\n") + 1,
        }

        return content

    def _extract_markdown_tables(self, text: str) -> List[TableData]:
        """Extract tables from Markdown text."""
        import re

        tables = []

        table_pattern = r"(?:[^\n]*\|[^\n]*\n)+(?::?-+:?.*\|.*\n)?(?:[^\n]*\|[^\n]*\n)*"
        matches = re.finditer(table_pattern, text)

        for table_idx, match in enumerate(matches):
            table_text = match.group().strip()
            if not table_text:
                continue

            lines = table_text.split("\n")
            rows_data = []

            for line in lines:
                if re.match(r"^[-:| ]+$", line):
                    continue

                cells = [c.strip() for c in line.strip("|").split("|")]
                if cells:
                    rows_data.append(cells)

            if len(rows_data) >= 2:
                table = TableData(
                    table_id=f"md_table_{table_idx + 1}",
                    page_number=1,
                )
                table.headers = rows_data[0]
                table.rows = rows_data[1:]

                try:
                    table.data_frame = pd.DataFrame(
                        table.rows,
                        columns=table.headers if table.headers else None,
                    )
                except Exception:
                    pass

                tables.append(table)

        return tables


class HTMLReader(DocumentReader):
    """Reader for HTML documents."""

    supported_formats = ["html", "htm"]

    def read(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = False,
        extract_links: bool = False,
    ) -> DocumentContent:
        """Read HTML document."""
        from bs4 import BeautifulSoup

        path = Path(file_path)
        file_stat = path.stat()

        content = DocumentContent(
            file_path=str(path),
            file_name=path.name,
            file_size=file_stat.st_size,
            document_type="html",
            page_count=1,
        )

        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                soup = BeautifulSoup(f.read(), "html.parser")

            for script in soup(["script", "style"]):
                script.decompose()

            content.text = soup.get_text(separator="\n", strip=True)

            tables = []
            if extract_tables:
                for table_idx, table_tag in enumerate(soup.find_all("table")):
                    rows_data = []
                    for row in table_tag.find_all("tr"):
                        cells = []
                        for cell in row.find_all(["th", "td"]):
                            cells.append(cell.get_text(strip=True))
                        if cells:
                            rows_data.append(cells)

                    if rows_data:
                        table = TableData(
                            table_id=f"html_table_{table_idx + 1}",
                            page_number=1,
                        )
                        if len(rows_data) > 1 and table_tag.find("th"):
                            table.headers = rows_data[0]
                            table.rows = rows_data[1:]
                        else:
                            table.rows = rows_data

                        try:
                            table.data_frame = pd.DataFrame(
                                table.rows,
                                columns=table.headers if table.headers else None,
                            )
                        except Exception:
                            pass

                        tables.append(table)

            content.tables = tables

            if extract_links:
                links = []
                for a_tag in soup.find_all("a", href=True):
                    links.append({
                        "text": a_tag.get_text(strip=True),
                        "href": a_tag["href"],
                        "title": a_tag.get("title", ""),
                    })
                content.links = links

            if extract_images:
                images = []
                for img_tag in soup.find_all("img", src=True):
                    images.append({
                        "src": img_tag["src"],
                        "alt": img_tag.get("alt", ""),
                        "title": img_tag.get("title", ""),
                    })
                content.images = images

            title_tag = soup.find("title")
            content.metadata = {
                "title": title_tag.get_text(strip=True) if title_tag else "",
            }

        except ImportError:
            raise ImportError(
                "beautifulsoup4 is required for HTML reading. "
                "Install it with: pip install beautifulsoup4"
            )

        return content


def get_reader_for_format(file_extension: str) -> Optional[type]:
    """Get the appropriate reader class for a file extension.

    Args:
        file_extension: File extension (with or without leading dot)

    Returns:
        DocumentReader subclass or None if no reader supports the format
    """
    readers = [PDFReader, WordReader, ExcelReader, TextReader, HTMLReader]

    for reader_cls in readers:
        if reader_cls.supports_format(file_extension):
            return reader_cls

    return None


def read_document(
    file_path: str,
    extract_tables: bool = True,
    extract_images: bool = False,
    extract_links: bool = False,
) -> DocumentContent:
    """Read a document using the appropriate reader.

    Args:
        file_path: Path to the document file
        extract_tables: Whether to extract tables
        extract_images: Whether to extract images
        extract_links: Whether to extract links

    Returns:
        DocumentContent object with extracted content

    Raises:
        ValueError: If no reader supports the file format
    """
    path = Path(file_path)
    ext = path.suffix

    reader_cls = get_reader_for_format(ext)

    if reader_cls is None:
        raise ValueError(f"No reader available for file format: {ext}")

    reader = reader_cls()
    return reader.read(
        file_path=file_path,
        extract_tables=extract_tables,
        extract_images=extract_images,
        extract_links=extract_links,
    )
