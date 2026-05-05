"""Document format converters.

This module provides document format conversion capabilities:
- PDF to Word
- Markdown to HTML
- HTML to Markdown
- Various text format conversions
"""

import json
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional


@dataclass
class ConversionResult:
    """Result of a document conversion."""

    success: bool
    source_path: str
    target_path: str
    source_format: str
    target_format: str
    error_message: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    converted_at: datetime = field(default_factory=datetime.now)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "success": self.success,
            "source_path": self.source_path,
            "target_path": self.target_path,
            "source_format": self.source_format,
            "target_format": self.target_format,
            "error_message": self.error_message,
            "metadata": self.metadata,
            "converted_at": self.converted_at.isoformat(),
        }


class DocumentConverter(ABC):
    """Abstract base class for document converters."""

    source_formats: List[str] = []
    target_formats: List[str] = []

    @abstractmethod
    def convert(
        self,
        source_path: str,
        target_path: str,
        **kwargs,
    ) -> ConversionResult:
        """Convert a document.

        Args:
            source_path: Path to source document
            target_path: Path to save converted document
            **kwargs: Additional conversion options

        Returns:
            ConversionResult
        """
        pass

    @classmethod
    def supports_conversion(cls, source_ext: str, target_ext: str) -> bool:
        """Check if this converter supports the conversion.

        Args:
            source_ext: Source file extension
            target_ext: Target file extension

        Returns:
            True if conversion is supported
        """
        source = source_ext.lower().lstrip(".")
        target = target_ext.lower().lstrip(".")
        return (
            source in [f.lower().lstrip(".") for f in cls.source_formats]
            and target in [f.lower().lstrip(".") for f in cls.target_formats]
        )


class MarkdownToHTMLConverter(DocumentConverter):
    """Convert Markdown to HTML."""

    source_formats = ["md", "markdown"]
    target_formats = ["html", "htm"]

    def convert(
        self,
        source_path: str,
        target_path: str,
        full_html: bool = True,
        title: str = "",
        css_url: Optional[str] = None,
        **kwargs,
    ) -> ConversionResult:
        """Convert Markdown to HTML.

        Args:
            source_path: Path to Markdown file
            target_path: Path to save HTML file
            full_html: Whether to create a full HTML document
            title: Document title (for full HTML)
            css_url: Optional CSS URL for styling
            **kwargs: Additional options
        """
        try:
            import markdown
        except ImportError:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="md",
                target_format="html",
                error_message="markdown library not installed. Install with: pip install markdown",
            )

        try:
            source = Path(source_path)
            target = Path(target_path)
            target.parent.mkdir(parents=True, exist_ok=True)

            with open(source, "r", encoding="utf-8") as f:
                md_content = f.read()

            md = markdown.Markdown(
                extensions=[
                    "extra",
                    "toc",
                    "tables",
                    "codehilite",
                    "fenced_code",
                    "meta",
                ]
            )
            html_body = md.convert(md_content)

            if full_html:
                css_link = ""
                if css_url:
                    css_link = f'<link rel="stylesheet" href="{css_url}">'

                doc_title = title or source.stem
                html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{doc_title}</title>
    {css_link}
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }}
        code {{
            background: #f4f4f4;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Fira Code', Consolas, monospace;
        }}
        pre {{
            background: #1f2937;
            color: #e5e7eb;
            padding: 16px;
            border-radius: 6px;
            overflow-x: auto;
        }}
        pre code {{
            background: none;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
        }}
        th, td {{
            border: 1px solid #e5e7eb;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background: #f9fafb;
            font-weight: 600;
        }}
        blockquote {{
            border-left: 4px solid #3b82f6;
            margin: 16px 0;
            padding-left: 16px;
            color: #6b7280;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        a {{
            color: #3b82f6;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        ul, ol {{
            padding-left: 24px;
        }}
        .toc {{
            background: #f9fafb;
            padding: 20px;
            border-radius: 8px;
            margin-bottom: 24px;
        }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""
            else:
                html_content = html_body

            with open(target, "w", encoding="utf-8") as f:
                f.write(html_content)

            return ConversionResult(
                success=True,
                source_path=str(source),
                target_path=str(target),
                source_format="md",
                target_format="html",
                metadata={
                    "full_html": full_html,
                    "word_count": len(md_content.split()),
                    "char_count": len(md_content),
                },
            )

        except Exception as e:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="md",
                target_format="html",
                error_message=str(e),
            )


class HTMLToMarkdownConverter(DocumentConverter):
    """Convert HTML to Markdown."""

    source_formats = ["html", "htm"]
    target_formats = ["md", "markdown"]

    def convert(
        self,
        source_path: str,
        target_path: str,
        **kwargs,
    ) -> ConversionResult:
        """Convert HTML to Markdown."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="html",
                target_format="md",
                error_message="beautifulsoup4 not installed. Install with: pip install beautifulsoup4",
            )

        try:
            source = Path(source_path)
            target = Path(target_path)
            target.parent.mkdir(parents=True, exist_ok=True)

            with open(source, "r", encoding="utf-8") as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, "html.parser")

            for script in soup(["script", "style"]):
                script.decompose()

            md_lines = []

            for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol", "table", "pre", "blockquote", "hr", "a", "img"]):
                if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    level = int(element.name[1])
                    text = element.get_text(strip=True)
                    if text:
                        md_lines.append(f"{'#' * level} {text}\n")

                elif element.name == "p":
                    text = element.get_text(strip=True)
                    if text:
                        md_lines.append(f"{text}\n\n")

                elif element.name == "ul":
                    for li in element.find_all("li", recursive=False):
                        text = li.get_text(strip=True)
                        if text:
                            md_lines.append(f"- {text}\n")
                    md_lines.append("\n")

                elif element.name == "ol":
                    for idx, li in enumerate(element.find_all("li", recursive=False), 1):
                        text = li.get_text(strip=True)
                        if text:
                            md_lines.append(f"{idx}. {text}\n")
                    md_lines.append("\n")

                elif element.name == "pre":
                    code = element.get_text()
                    md_lines.append(f"```\n{code}\n```\n\n")

                elif element.name == "blockquote":
                    text = element.get_text(strip=True)
                    if text:
                        for line in text.split("\n"):
                            if line.strip():
                                md_lines.append(f"> {line}\n")
                        md_lines.append("\n")

                elif element.name == "hr":
                    md_lines.append("---\n\n")

                elif element.name == "table":
                    rows = []
                    headers = []

                    for tr in element.find_all("tr"):
                        row_cells = []
                        for cell in tr.find_all(["th", "td"]):
                            row_cells.append(cell.get_text(strip=True))
                        if any(row_cells):
                            rows.append(row_cells)
                            if tr.find("th") and not headers:
                                headers = row_cells

                    if rows:
                        if headers:
                            md_lines.append(f"| {' | '.join(headers)} |\n")
                            md_lines.append(f"| {' | '.join(['---' for _ in headers])} |\n")
                            data_rows = rows[1:] if len(rows) > 1 else []
                        else:
                            if rows:
                                md_lines.append(f"| {' | '.join(['col' + str(i+1) for i in range(len(rows[0]))])} |\n")
                                md_lines.append(f"| {' | '.join(['---' for _ in rows[0]])} |\n")
                            data_rows = rows

                        for row in data_rows:
                            md_lines.append(f"| {' | '.join(row)} |\n")
                        md_lines.append("\n")

            md_content = "".join(md_lines)

            with open(target, "w", encoding="utf-8") as f:
                f.write(md_content)

            return ConversionResult(
                success=True,
                source_path=str(source),
                target_path=str(target),
                source_format="html",
                target_format="md",
                metadata={
                    "line_count": md_content.count("\n") + 1,
                    "char_count": len(md_content),
                },
            )

        except Exception as e:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="html",
                target_format="md",
                error_message=str(e),
            )


class PDFToWordConverter(DocumentConverter):
    """Convert PDF to Word document.

    Note: This is a text-based conversion. For more advanced PDF conversion,
    consider using additional libraries like pdf2docx.
    """

    source_formats = ["pdf"]
    target_formats = ["docx"]

    def convert(
        self,
        source_path: str,
        target_path: str,
        include_tables: bool = True,
        **kwargs,
    ) -> ConversionResult:
        """Convert PDF to Word.

        Args:
            source_path: Path to PDF file
            target_path: Path to save DOCX file
            include_tables: Whether to include extracted tables
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="pdf",
                target_format="docx",
                error_message="python-docx not installed. Install with: pip install python-docx",
            )

        try:
            import pdfplumber
        except ImportError:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="pdf",
                target_format="docx",
                error_message="pdfplumber not installed. Install with: pip install pdfplumber",
            )

        try:
            source = Path(source_path)
            target = Path(target_path)
            target.parent.mkdir(parents=True, exist_ok=True)

            doc = Document()

            title = doc.add_heading(f"Converted from: {source.name}", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            with pdfplumber.open(source) as pdf:
                total_pages = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    doc.add_heading(f"Page {page_num} of {total_pages}", level=1)

                    text = page.extract_text()
                    if text:
                        for para_text in text.split("\n\n"):
                            if para_text.strip():
                                para = doc.add_paragraph(para_text.strip())
                                para.paragraph_format.space_after = Pt(6)

                    if include_tables:
                        tables = page.extract_tables()
                        for table_idx, table_data in enumerate(tables):
                            if table_data:
                                doc.add_paragraph()
                                doc.add_heading(f"Table {table_idx + 1}", level=3)

                                if table_data:
                                    rows = len(table_data)
                                    cols = len(table_data[0]) if table_data else 0

                                    table = doc.add_table(rows=rows, cols=cols)
                                    table.style = "Table Grid"

                                    for row_idx, row_data in enumerate(table_data):
                                        row_cells = table.rows[row_idx].cells
                                        for col_idx, cell_text in enumerate(row_data):
                                            if col_idx < len(row_cells):
                                                row_cells[col_idx].text = str(cell_text) if cell_text else ""

                    if page_num < total_pages:
                        doc.add_page_break()

            doc.save(target)

            return ConversionResult(
                success=True,
                source_path=str(source),
                target_path=str(target),
                source_format="pdf",
                target_format="docx",
                metadata={
                    "page_count": total_pages,
                    "include_tables": include_tables,
                },
            )

        except Exception as e:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format="pdf",
                target_format="docx",
                error_message=str(e),
            )


class TextFormatConverter:
    """Simple text format converter."""

    CONVERTERS = [
        MarkdownToHTMLConverter,
        HTMLToMarkdownConverter,
        PDFToWordConverter,
    ]

    @classmethod
    def get_converter(
        cls,
        source_format: str,
        target_format: str,
    ) -> Optional[type]:
        """Get the appropriate converter for a format pair.

        Args:
            source_format: Source file extension
            target_format: Target file extension

        Returns:
            Converter class or None
        """
        for converter_cls in cls.CONVERTERS:
            if converter_cls.supports_conversion(source_format, target_format):
                return converter_cls
        return None

    @classmethod
    def convert(
        cls,
        source_path: str,
        target_path: str,
        **kwargs,
    ) -> ConversionResult:
        """Convert a document based on file extensions.

        Args:
            source_path: Source file path
            target_path: Target file path
            **kwargs: Additional conversion options

        Returns:
            ConversionResult
        """
        source = Path(source_path)
        target = Path(target_path)

        source_ext = source.suffix
        target_ext = target.suffix

        converter_cls = cls.get_converter(source_ext, target_ext)

        if converter_cls is None:
            return ConversionResult(
                success=False,
                source_path=source_path,
                target_path=target_path,
                source_format=source_ext,
                target_format=target_ext,
                error_message=f"No converter available for {source_ext} -> {target_ext}",
            )

        converter = converter_cls()
        return converter.convert(source_path, target_path, **kwargs)

    @classmethod
    def list_supported_conversions(cls) -> List[Dict[str, str]]:
        """List all supported conversions.

        Returns:
            List of dictionaries with source and target formats
        """
        conversions = []
        for converter_cls in cls.CONVERTERS:
            for source in converter_cls.source_formats:
                for target in converter_cls.target_formats:
                    conversions.append({
                        "source": source,
                        "target": target,
                        "converter": converter_cls.__name__,
                    })
        return conversions
