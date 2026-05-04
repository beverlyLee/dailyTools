from typing import Dict, Any, Optional, List
from datetime import datetime
import io

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class WordDocumentService:
    def __init__(self):
        self.default_font = "宋体"
        self.default_font_size = None
        if HAS_DOCX:
            from docx.shared import Pt
            self.default_font_size = Pt(14)

    def _set_cell_shading(self, cell, color: str):
        if not HAS_DOCX:
            return
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _add_heading(
        self,
        doc: Any,
        text: str,
        level: int = 0,
        align: str = "center"
    ):
        if not HAS_DOCX:
            return
        para = doc.add_paragraph()
        run = para.add_run(text)
        
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if level == 0:
            run.font.size = Pt(22)
            run.font.bold = True
            run.font.name = self.default_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
            para.space_before = Pt(24)
            para.space_after = Pt(24)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = self.default_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
            para.space_before = Pt(18)
            para.space_after = Pt(12)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.name = self.default_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
            para.space_before = Pt(12)
            para.space_after = Pt(6)

    def _add_paragraph(
        self,
        doc: Any,
        text: str,
        indent: bool = True,
        bold: bool = False,
        size: Optional[Any] = None,
        align: str = "left"
    ):
        if not HAS_DOCX:
            return
        para = doc.add_paragraph()
        
        if align == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if indent:
            para.first_line_indent = Inches(0.5)
        
        para.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.line_spacing = 1.5
        para.space_after = Pt(6)
        
        run = para.add_run(text)
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = size or self.default_font_size
        run.font.bold = bold

    def _add_table(
        self,
        doc: Any,
        rows: int,
        cols: int,
        data: List[List[str]],
        header_row: bool = True
    ):
        if not HAS_DOCX:
            return
        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = self.default_font
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
                        run.font.size = Pt(12)
                
                if i == 0 and header_row:
                    self._set_cell_shading(cell, "D9E2F3")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True

    def generate_complaint(self, context: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            raise ImportError("python-docx 模块未安装，请安装: pip install python-docx")
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._add_heading(doc, "民事起诉状", level=0, align="center")

        court_name = context.get("court_name", "______人民法院")
        self._add_paragraph(doc, f"原告：{court_name}", indent=False, align="left")
        doc.add_paragraph()

        plaintiff_name = context.get("plaintiff_name", "______")
        plaintiff_gender = context.get("plaintiff_gender", "")
        plaintiff_ethnicity = context.get("plaintiff_ethnicity", "")
        plaintiff_occupation = context.get("plaintiff_occupation", "")
        plaintiff_address = context.get("plaintiff_address", "")
        plaintiff_phone = context.get("plaintiff_phone", "")
        
        plaintiff_text = f"原告：{plaintiff_name}，{plaintiff_gender}，{plaintiff_ethnicity}"
        if plaintiff_occupation:
            plaintiff_text += f"，{plaintiff_occupation}"
        self._add_paragraph(doc, plaintiff_text, indent=False)
        
        if plaintiff_address:
            self._add_paragraph(doc, f"住址：{plaintiff_address}", indent=False)
        if plaintiff_phone:
            self._add_paragraph(doc, f"联系电话：{plaintiff_phone}", indent=False)

        doc.add_paragraph()

        defendant_name = context.get("defendant_name", "______")
        defendant_gender = context.get("defendant_gender", "")
        defendant_ethnicity = context.get("defendant_ethnicity", "")
        defendant_occupation = context.get("defendant_occupation", "")
        defendant_address = context.get("defendant_address", "")
        defendant_phone = context.get("defendant_phone", "")
        
        defendant_text = f"被告：{defendant_name}，{defendant_gender}，{defendant_ethnicity}"
        if defendant_occupation:
            defendant_text += f"，{defendant_occupation}"
        self._add_paragraph(doc, defendant_text, indent=False)
        
        if defendant_address:
            self._add_paragraph(doc, f"住址：{defendant_address}", indent=False)
        if defendant_phone:
            self._add_paragraph(doc, f"联系电话：{defendant_phone}", indent=False)

        doc.add_paragraph()

        self._add_heading(doc, "诉讼请求", level=1, align="left")
        
        claims = context.get("claims", [])
        if claims:
            for i, claim in enumerate(claims, 1):
                self._add_paragraph(doc, f"{i}. {claim}")
        else:
            self._add_paragraph(doc, "1. 请求法院判令被告________；")
            self._add_paragraph(doc, "2. 请求法院判令被告承担本案诉讼费用。")

        self._add_heading(doc, "事实与理由", level=1, align="left")
        
        facts = context.get("facts_and_reasons", "")
        if facts:
            paragraphs = facts.split("\n")
            for para in paragraphs:
                if para.strip():
                    self._add_paragraph(doc, para.strip())
        else:
            self._add_paragraph(doc, "原告与被告于______年______月______日因________________发生纠纷。")
            self._add_paragraph(doc, "综上所述，原告认为，被告的行为严重侵害了原告的合法权益。为维护原告的合法权益，根据《中华人民共和国民事诉讼法》及相关法律规定，特向贵院提起诉讼，恳请依法判决。")

        self._add_heading(doc, "证据和证据来源，证人姓名和住所", level=1, align="left")
        
        evidence_list = context.get("evidence_list", [])
        if evidence_list:
            for i, evidence in enumerate(evidence_list, 1):
                name = evidence.get("name", "") if isinstance(evidence, dict) else ""
                desc = evidence.get("description", "") if isinstance(evidence, dict) else ""
                if desc:
                    self._add_paragraph(doc, f"{i}. {name}：{desc}")
                else:
                    self._add_paragraph(doc, f"{i}. {name}" if name else f"{i}. {evidence}")
        else:
            self._add_paragraph(doc, "1. ________________（证据名称），证明________________；")
            self._add_paragraph(doc, "2. 其他相关证据材料。")

        doc.add_paragraph()

        court_name = context.get("court_name", "______人民法院")
        self._add_paragraph(doc, f"此致", indent=False)
        self._add_paragraph(doc, court_name, indent=False)

        doc.add_paragraph()
        doc.add_paragraph()

        today = datetime.now().strftime("%Y年%m月%d日")
        filing_date = context.get("filing_date", today)
        
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"起诉人：{plaintiff_name}")
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(filing_date)
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_contract(self, context: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            raise ImportError("python-docx 模块未安装，请安装: pip install python-docx")
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._add_heading(doc, "借 款 合 同", level=0, align="center")

        contract_no = context.get("contract_no", "")
        if contract_no:
            self._add_paragraph(doc, f"合同编号：{contract_no}", indent=False, align="right")

        doc.add_paragraph()

        party_a_name = context.get("party_a_name", "甲方（出借人）")
        party_a_address = context.get("party_a_address", "")
        party_a_contact = context.get("party_a_contact", "")
        
        self._add_paragraph(doc, f"甲方（出借人）：{party_a_name}", indent=False)
        if party_a_address:
            self._add_paragraph(doc, f"住址：{party_a_address}", indent=False)
        if party_a_contact:
            self._add_paragraph(doc, f"联系电话：{party_a_contact}", indent=False)

        doc.add_paragraph()

        party_b_name = context.get("party_b_name", "乙方（借款人）")
        party_b_address = context.get("party_b_address", "")
        party_b_contact = context.get("party_b_contact", "")
        
        self._add_paragraph(doc, f"乙方（借款人）：{party_b_name}", indent=False)
        if party_b_address:
            self._add_paragraph(doc, f"住址：{party_b_address}", indent=False)
        if party_b_contact:
            self._add_paragraph(doc, f"联系电话：{party_b_contact}", indent=False)

        doc.add_paragraph()

        self._add_paragraph(doc, "甲乙双方本着平等自愿、诚实信用的原则，经协商一致，达成本合同，并保证共同遵守执行。")

        self._add_heading(doc, "第一条 借款金额", level=2, align="left")
        principal = context.get("principal_amount", "")
        if principal:
            try:
                if '.' in str(principal):
                    principal_int = int(float(principal))
                else:
                    principal_int = int(principal)
                principal_cn = self._num_to_cn(principal_int)
                self._add_paragraph(doc, f"甲方向乙方出借人民币（大写）{principal_cn}整（小写：￥{principal}元）。")
            except (ValueError, TypeError):
                self._add_paragraph(doc, "甲方向乙方出借人民币（大写）________________整（小写：￥________元）。")
        else:
            self._add_paragraph(doc, "甲方向乙方出借人民币（大写）________________整（小写：￥________元）。")

        self._add_heading(doc, "第二条 借款用途", level=2, align="left")
        purpose = context.get("purpose", "")
        if purpose:
            self._add_paragraph(doc, f"乙方借款用于：{purpose}。")
        else:
            self._add_paragraph(doc, "乙方借款用于：________________。")

        self._add_heading(doc, "第三条 借款期限", level=2, align="left")
        term = context.get("borrowing_term", "")
        if term:
            self._add_paragraph(doc, f"借款期限为：{term}。")
        else:
            self._add_paragraph(doc, "借款期限为：自______年______月______日起至______年______月______日止。")

        self._add_heading(doc, "第四条 借款利率", level=2, align="left")
        rate = context.get("interest_rate", "")
        if rate:
            self._add_paragraph(doc, f"借款利率为：{rate}。")
        else:
            self._add_paragraph(doc, f"借款利率为：月利率______%（年利率______%）。")

        self._add_heading(doc, "第五条 还款方式", level=2, align="left")
        method = context.get("repayment_method", "")
        if method:
            self._add_paragraph(doc, f"还款方式：{method}。")
        else:
            self._add_paragraph(doc, "还款方式：乙方应于借款到期日一次性偿还本金及利息。")

        self._add_heading(doc, "第六条 违约责任", level=2, align="left")
        self._add_paragraph(doc, "1. 乙方如逾期还款，应按逾期金额的每日万分之五支付违约金。")
        self._add_paragraph(doc, "2. 乙方如逾期超过30天，甲方有权解除本合同，并要求乙方提前偿还全部借款本息。")

        self._add_heading(doc, "第七条 争议解决", level=2, align="left")
        self._add_paragraph(doc, "本合同履行过程中发生的争议，由双方协商解决；协商不成的，可向合同签订地人民法院提起诉讼。")

        doc.add_paragraph()
        self._add_paragraph(doc, "本合同自双方签字盖章之日起生效。本合同一式两份，甲乙双方各执一份，具有同等法律效力。", indent=False)

        doc.add_paragraph()
        doc.add_paragraph()

        signing_date = context.get("signing_date", datetime.now().strftime("%Y年%m月%d日"))
        signing_place = context.get("signing_place", "")

        table_data = [
            ["甲方（出借人）", "乙方（借款人）"],
            [f"签字：{party_a_name}", f"签字：{party_b_name}"],
            [f"日期：{signing_date}", f"日期：{signing_date}"],
        ]
        if signing_place:
            table_data.append([f"签订地点：{signing_place}", f"签订地点：{signing_place}"])

        self._add_table(doc, len(table_data), 2, table_data, header_row=False)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_claim_letter(self, context: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            raise ImportError("python-docx 模块未安装，请安装: pip install python-docx")
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._add_heading(doc, "索 赔 函", level=0, align="center")

        claimant_name = context.get("claimant_name", "索赔方")
        claimant_address = context.get("claimant_address", "")
        claimant_contact = context.get("claimant_contact", "")
        letter_date = context.get("claim_date", datetime.now().strftime("%Y年%m月%d日"))
        
        self._add_paragraph(doc, f"致：{context.get('respondent_name', '被索赔方')}", indent=False)
        if context.get("respondent_address"):
            self._add_paragraph(doc, f"地址：{context.get('respondent_address', '')}", indent=False)

        doc.add_paragraph()

        self._add_paragraph(doc, "事由：关于____________事项的索赔", bold=True)

        doc.add_paragraph()

        self._add_paragraph(doc, f"贵司/阁下于______年______月______日因________________，给我司/本人造成了经济损失。根据双方约定/相关法律规定，特向贵司/阁下提出如下索赔：")

        claim_amount = context.get("claim_amount", "")
        if claim_amount:
            try:
                if '.' in str(claim_amount):
                    amount_int = int(float(claim_amount))
                else:
                    amount_int = int(claim_amount)
                amount_cn = self._num_to_cn(amount_int)
                self._add_paragraph(doc, f"一、索赔金额：人民币（大写）{amount_cn}整（小写：￥{claim_amount}元）。")
            except (ValueError, TypeError):
                self._add_paragraph(doc, "一、索赔金额：人民币（大写）________________整（小写：￥________元）。")
        else:
            self._add_paragraph(doc, "一、索赔金额：人民币（大写）________________整（小写：￥________元）。")

        self._add_paragraph(doc, "二、索赔依据：")
        claim_reason = context.get("claim_reason", "")
        if claim_reason:
            paragraphs = claim_reason.split("\n")
            for para in paragraphs:
                if para.strip():
                    self._add_paragraph(doc, para.strip())
        else:
            self._add_paragraph(doc, "1. 双方于______年______月______日签订的《______合同》；")
            self._add_paragraph(doc, "2. 相关损失证明材料；")
            self._add_paragraph(doc, "3. 其他相关证据。")

        deadline = context.get("deadline_date", "")
        if deadline:
            self._add_paragraph(doc, f"请贵司/阁下于{deadline}前将上述款项支付至以下账户：")
        else:
            self._add_paragraph(doc, "请贵司/阁下于______年______月______日前将上述款项支付至以下账户：")

        self._add_paragraph(doc, "户名：________________")
        self._add_paragraph(doc, "账号：________________")
        self._add_paragraph(doc, "开户行：________________")

        self._add_paragraph(doc, "如贵司/阁下未在上述期限内履行赔偿义务，我司/本人将通过法律途径维护自身合法权益，届时产生的诉讼费、律师费、差旅费等均由贵司/阁下承担。")

        self._add_paragraph(doc, "特此函告！")

        doc.add_paragraph()
        doc.add_paragraph()

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"索赔人：{claimant_name}")
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(letter_date)
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        attachments = context.get("evidence_attachments", [])
        if attachments:
            doc.add_paragraph()
            self._add_paragraph(doc, "附件：", bold=True)
            for i, att in enumerate(attachments, 1):
                self._add_paragraph(doc, f"{i}. {att}")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_power_of_attorney(self, context: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            raise ImportError("python-docx 模块未安装，请安装: pip install python-docx")
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

        self._add_heading(doc, "授 权 委 托 书", level=0, align="center")

        principal_name = context.get("principal_name", "委托人")
        principal_gender = context.get("principal_gender", "")
        principal_ethnicity = context.get("principal_ethnicity", "")
        principal_id = context.get("principal_id_number", "")
        principal_address = context.get("principal_address", "")

        self._add_paragraph(doc, f"委托人：{principal_name}，{principal_gender}，{principal_ethnicity}", indent=False)
        if principal_id:
            self._add_paragraph(doc, f"身份证号码：{principal_id}", indent=False)
        if principal_address:
            self._add_paragraph(doc, f"住址：{principal_address}", indent=False)

        doc.add_paragraph()

        agent_name = context.get("agent_name", "受托人")
        agent_gender = context.get("agent_gender", "")
        agent_ethnicity = context.get("agent_ethnicity", "")
        agent_id = context.get("agent_id_number", "")
        agent_address = context.get("agent_address", "")

        self._add_paragraph(doc, f"受托人：{agent_name}，{agent_gender}，{agent_ethnicity}", indent=False)
        if agent_id:
            self._add_paragraph(doc, f"身份证号码：{agent_id}", indent=False)
        if agent_address:
            self._add_paragraph(doc, f"住址：{agent_address}", indent=False)

        doc.add_paragraph()

        purpose = context.get("authorization_purpose", "")
        if purpose:
            self._add_paragraph(doc, f"现委托{agent_name}在我与________________一案中，作为我的委托代理人。")
        else:
            self._add_paragraph(doc, f"现委托{agent_name}在我与________________一案中，作为我的委托代理人。")

        self._add_heading(doc, "委托权限", level=1, align="left")

        scope = context.get("authorization_scope", "")
        if scope:
            self._add_paragraph(doc, scope)
        else:
            self._add_paragraph(doc, "一般授权：代为提交诉讼材料、代为出庭、代为陈述案情、代为举证、代为质证、代为辩论、代为签收法律文书。")
            self._add_paragraph(doc, "特别授权（如需要请勾选）：□ 代为承认、放弃、变更诉讼请求；□ 代为进行和解；□ 代为提起反诉或上诉。")

        self._add_heading(doc, "委托期限", level=1, align="left")

        start_date = context.get("authorization_start_date", "")
        end_date = context.get("authorization_end_date", "")
        
        if start_date and end_date:
            self._add_paragraph(doc, f"自{start_date}起至{end_date}止。")
        else:
            self._add_paragraph(doc, "自本委托书签署之日起至本案终结止（包括一审、二审、执行阶段）。")

        self._add_paragraph(doc, "受托人在上述委托权限范围内所签署的一切文件，委托人均予以承认，并承担由此产生的一切法律后果。")
        self._add_paragraph(doc, "受托人无转委托权。")

        doc.add_paragraph()
        doc.add_paragraph()

        issuing_date = context.get("issuing_date", datetime.now().strftime("%Y年%m月%d日"))

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"委托人：{principal_name}")
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(issuing_date)
        run.font.name = self.default_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.default_font)
        run.font.size = self.default_font_size

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _num_to_cn(self, num: int) -> str:
        cn_nums = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
        cn_units = ["", "拾", "佰", "仟", "万", "拾", "佰", "仟", "亿"]
        
        if num == 0:
            return "零"
        
        result = ""
        num_str = str(num)
        length = len(num_str)
        
        for i, digit in enumerate(num_str):
            d = int(digit)
            unit_index = length - i - 1
            if d != 0:
                result += cn_nums[d] + cn_units[unit_index]
            else:
                if unit_index % 4 == 0 and unit_index != 0:
                    result += cn_units[unit_index]
                elif len(result) > 0 and not result.endswith("零"):
                    result += "零"
        
        result = result.replace("零万", "万")
        result = result.replace("零亿", "亿")
        result = result.rstrip("零")
        
        return result

    def generate_document(
        self,
        template_name: str,
        context: Dict[str, Any]
    ) -> bytes:
        template_map = {
            "complaint_template.jinja2": self.generate_complaint,
            "contract_template.jinja2": self.generate_contract,
            "claim_letter_template.jinja2": self.generate_claim_letter,
            "power_of_attorney_template.jinja2": self.generate_power_of_attorney,
        }
        
        generator = template_map.get(template_name)
        if generator:
            return generator(context)
        
        return self._generate_generic_document(context)

    def _generate_generic_document(self, context: Dict[str, Any]) -> bytes:
        if not HAS_DOCX:
            raise ImportError("python-docx 模块未安装，请安装: pip install python-docx")
        
        doc = Document()
        
        title = context.get("title", "法律文书")
        self._add_heading(doc, title, level=0, align="center")
        
        content = context.get("content", "")
        if content:
            paragraphs = content.split("\n")
            for para in paragraphs:
                if para.strip():
                    self._add_paragraph(doc, para.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


word_document_service = WordDocumentService()
