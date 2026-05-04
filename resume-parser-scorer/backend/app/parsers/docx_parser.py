from docx import Document
from docx.table import Table
import re
from typing import Optional


class DOCXParser:
    """Word文档简历解析器"""
    
    def __init__(self):
        pass
    
    def parse(self, file_path: str) -> str:
        """
        解析Word文档，提取文本内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            提取的文本内容
        """
        text = ""
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    text += table_text + "\n"
            
            return self._clean_text(text)
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")
    
    def _extract_table_text(self, table: Table) -> str:
        """
        提取表格文本内容
        
        Args:
            table: Word表格对象
            
        Returns:
            表格文本内容
        """
        table_text = ""
        for row in table.rows:
            row_text = ""
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text += cell_text + " "
            if row_text.strip():
                table_text += row_text.strip() + "\n"
        return table_text
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本内容
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        text = text.strip()
        return text
    
    def extract_by_regex(self, text: str, pattern: str) -> Optional[str]:
        """
        使用正则表达式提取文本
        
        Args:
            text: 文本内容
            pattern: 正则表达式模式
            
        Returns:
            匹配的内容，未匹配返回None
        """
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if match:
            return match.group(1).strip() if match.groups() else match.group().strip()
        return None
