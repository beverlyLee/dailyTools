import pytest
import tempfile
import os
from docx import Document
import pdfplumber

from app.parsers.pdf_parser import PDFParser
from app.parsers.docx_parser import DOCXParser


class TestPDFParser:
    """PDF解析器测试"""
    
    def setup_method(self):
        self.parser = PDFParser()
    
    def test_init(self):
        """测试初始化"""
        assert self.parser is not None
    
    def test_clean_text(self):
        """测试文本清理"""
        dirty_text = "  这是   一些  \n\n\n  测试文本  \n"
        cleaned = self.parser._clean_text(dirty_text)
        assert "\n\n" not in cleaned
        assert "  " not in cleaned
        assert cleaned.strip() == "这是 一些 \n 测试文本"
    
    def test_extract_by_regex(self):
        """测试正则提取"""
        text = "姓名: 张三\n电话: 13800138000\n邮箱: zhangsan@example.com"
        
        name = self.parser.extract_by_regex(text, r'姓名[:：]\s*([\u4e00-\u9fa5]+)')
        assert name == "张三"
        
        phone = self.parser.extract_by_regex(text, r'1[3-9]\d{9}')
        assert phone == "13800138000"


class TestDOCXParser:
    """Word文档解析器测试"""
    
    def setup_method(self):
        self.parser = DOCXParser()
    
    def test_init(self):
        """测试初始化"""
        assert self.parser is not None
    
    def test_clean_text(self):
        """测试文本清理"""
        dirty_text = "  这是   一些  \n\n\n  测试文本  \n"
        cleaned = self.parser._clean_text(dirty_text)
        assert "\n\n" not in cleaned
        assert "  " not in cleaned
    
    def test_extract_by_regex(self):
        """测试正则提取"""
        text = "姓名: 李四\n电话: 13900139000\n邮箱: lisi@example.com"
        
        name = self.parser.extract_by_regex(text, r'姓名[:：]\s*([\u4e00-\u9fa5]+)')
        assert name == "李四"


class TestDocumentParsing:
    """文档解析集成测试"""
    
    def test_docx_creation_and_parsing(self):
        """测试创建并解析Word文档"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            doc = Document()
            doc.add_heading('张三', 0)
            doc.add_paragraph('姓名: 张三')
            doc.add_paragraph('性别: 男')
            doc.add_paragraph('电话: 13800138000')
            doc.add_paragraph('邮箱: zhangsan@example.com')
            doc.add_paragraph('学历: 本科')
            doc.add_paragraph('专业: 计算机科学与技术')
            doc.save(tmp_path)
            
            parser = DOCXParser()
            text = parser.parse(tmp_path)
            
            assert '张三' in text
            assert '13800138000' in text
            assert 'zhangsan@example.com' in text
            assert '本科' in text
            assert '计算机科学与技术' in text
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_text_extraction(self):
        """测试文本提取功能"""
        parser = PDFParser()
        
        test_text = """
        张三
        高级软件工程师
        
        联系方式
        电话: 13800138000
        邮箱: zhangsan@example.com
        
        教育背景
        北京大学 - 计算机科学与技术 - 本科
        2016年09月 至 2020年06月
        
        工作经历
        阿里巴巴 - 高级软件工程师
        2020年07月 至今
        负责后端系统开发，使用Java、Python等技术栈
        """
        
        cleaned = parser._clean_text(test_text)
        assert '张三' in cleaned
        assert '13800138000' in cleaned
        assert '阿里巴巴' in cleaned
        assert 'Java' in cleaned
        assert 'Python' in cleaned


class TestParserEdgeCases:
    """解析器边界情况测试"""
    
    def test_empty_text_cleaning(self):
        """测试空文本清理"""
        parser = PDFParser()
        
        assert parser._clean_text("") == ""
        assert parser._clean_text("   ") == ""
        assert parser._clean_text("\n\n\n") == ""
    
    def test_regex_no_match(self):
        """测试正则不匹配情况"""
        parser = PDFParser()
        
        result = parser.extract_by_regex("普通文本", r'不存在的模式')
        assert result is None
    
    def test_special_characters(self):
        """测试特殊字符处理"""
        parser = PDFParser()
        
        special_text = "姓名: 张\n三\t电话: 138\t0013 8000"
        cleaned = parser._clean_text(special_text)
        assert '张' in cleaned
        assert '138' in cleaned
