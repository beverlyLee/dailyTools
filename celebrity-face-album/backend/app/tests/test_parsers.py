import pytest
import os
import tempfile
from unittest.mock import Mock, patch, MagicMock
from docx import Document
from docx.shared import Pt


class TestPDFParser:
    """PDF解析器测试"""
    
    def setup_method(self):
        """测试前准备"""
        from ..parsers.pdf_parser import PDFParser
        self.parser = PDFParser()
    
    def test_clean_text(self):
        """测试文本清理功能"""
        # 测试多个空格替换
        dirty_text = "这是   一个   测试   文本"
        cleaned = self.parser._clean_text(dirty_text)
        assert "   " not in cleaned
        assert "  " not in cleaned
        
        # 测试多个换行替换
        dirty_text = "第一行\n\n\n第二行\n\n第三行"
        cleaned = self.parser._clean_text(dirty_text)
        assert "\n\n" not in cleaned
        
        # 测试首尾空白
        dirty_text = "   测试文本   "
        cleaned = self.parser._clean_text(dirty_text)
        assert cleaned == "测试文本"
    
    def test_extract_by_regex(self):
        """测试正则表达式提取"""
        text = "姓名：张三\n电话：13800138000\n邮箱：zhangsan@example.com"
        
        # 测试提取姓名
        name = self.parser.extract_by_regex(text, r'姓名[:：]\s*(\S+)')
        assert name == "张三"
        
        # 测试提取电话
        phone = self.parser.extract_by_regex(text, r'电话[:：]\s*(\d+)')
        assert phone == "13800138000"
        
        # 测试提取邮箱
        email = self.parser.extract_by_regex(text, r'邮箱[:：]\s*(\S+)')
        assert email == "zhangsan@example.com"
        
        # 测试不匹配的情况
        result = self.parser.extract_by_regex(text, r'不存在的字段[:：]\s*(\S+)')
        assert result is None
    
    @patch('pdfplumber.open')
    def test_parse_success(self, mock_pdfplumber_open):
        """测试成功解析PDF文件"""
        # 创建模拟的PDF页面
        mock_page = Mock()
        mock_page.extract_text.return_value = "这是PDF文本内容\n第二行内容"
        
        mock_pdf = Mock()
        mock_pdf.pages = [mock_page]
        mock_pdfplumber_open.return_value.__enter__.return_value = mock_pdf
        
        # 创建临时文件路径
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            result = self.parser.parse(temp_path)
            
            # 验证结果
            assert "这是PDF文本内容" in result
            assert "第二行内容" in result
            
            # 验证pdfplumber.open被调用
            mock_pdfplumber_open.assert_called_once_with(temp_path)
        finally:
            os.unlink(temp_path)
    
    @patch('pdfplumber.open')
    def test_parse_with_exception(self, mock_pdfplumber_open):
        """测试解析PDF文件时抛出异常"""
        mock_pdfplumber_open.side_effect = Exception("PDF文件损坏")
        
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            with pytest.raises(Exception) as excinfo:
                self.parser.parse(temp_path)
            
            assert "PDF解析失败" in str(excinfo.value)
            assert "PDF文件损坏" in str(excinfo.value)
        finally:
            os.unlink(temp_path)


class TestDOCXParser:
    """Word文档解析器测试"""
    
    def setup_method(self):
        """测试前准备"""
        from ..parsers.docx_parser import DOCXParser
        self.parser = DOCXParser()
    
    def test_clean_text(self):
        """测试文本清理功能"""
        # 测试多个空格替换
        dirty_text = "这是   一个   测试   文本"
        cleaned = self.parser._clean_text(dirty_text)
        assert "   " not in cleaned
        assert "  " not in cleaned
        
        # 测试多个换行替换
        dirty_text = "第一行\n\n\n第二行\n\n第三行"
        cleaned = self.parser._clean_text(dirty_text)
        assert "\n\n" not in cleaned
        
        # 测试首尾空白
        dirty_text = "   测试文本   "
        cleaned = self.parser._clean_text(dirty_text)
        assert cleaned == "测试文本"
    
    def test_extract_by_regex(self):
        """测试正则表达式提取"""
        text = "姓名：李四\n电话：13900139000\n邮箱：lisi@example.com"
        
        # 测试提取姓名
        name = self.parser.extract_by_regex(text, r'姓名[:：]\s*(\S+)')
        assert name == "李四"
        
        # 测试提取电话
        phone = self.parser.extract_by_regex(text, r'电话[:：]\s*(\d+)')
        assert phone == "13900139000"
        
        # 测试不匹配的情况
        result = self.parser.extract_by_regex(text, r'不存在的字段[:：]\s*(\S+)')
        assert result is None
    
    def test_parse_with_real_docx(self):
        """测试解析真实的Word文档"""
        # 创建一个临时的Word文档
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            # 创建Word文档
            doc = Document()
            
            # 添加段落
            doc.add_paragraph("这是测试文档的第一段落")
            doc.add_paragraph("第二段落包含更多内容")
            
            # 添加表格
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "单元格1-1"
            table.cell(0, 1).text = "单元格1-2"
            table.cell(1, 0).text = "单元格2-1"
            table.cell(1, 1).text = "单元格2-2"
            
            doc.save(temp_path)
            
            # 解析文档
            result = self.parser.parse(temp_path)
            
            # 验证段落文本被提取
            assert "这是测试文档的第一段落" in result
            assert "第二段落包含更多内容" in result
            
            # 验证表格文本被提取
            assert "单元格1-1" in result
            assert "单元格1-2" in result
        finally:
            os.unlink(temp_path)
    
    @patch('docx.Document')
    def test_parse_with_exception(self, mock_document):
        """测试解析Word文档时抛出异常"""
        mock_document.side_effect = Exception("Word文件损坏")
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            with pytest.raises(Exception) as excinfo:
                self.parser.parse(temp_path)
            
            assert "Word文档解析失败" in str(excinfo.value)
            assert "Word文件损坏" in str(excinfo.value)
        finally:
            os.unlink(temp_path)


class TestIntegration:
    """集成测试"""
    
    def test_pdf_and_docx_parsers_same_interface(self):
        """测试PDF和DOCX解析器有相同的接口"""
        from ..parsers.pdf_parser import PDFParser
        from ..parsers.docx_parser import DOCXParser
        
        pdf_parser = PDFParser()
        docx_parser = DOCXParser()
        
        # 验证两个解析器都有parse方法
        assert hasattr(pdf_parser, 'parse')
        assert hasattr(docx_parser, 'parse')
        
        # 验证两个解析器都有_clean_text方法
        assert hasattr(pdf_parser, '_clean_text')
        assert hasattr(docx_parser, '_clean_text')
        
        # 验证两个解析器都有extract_by_regex方法
        assert hasattr(pdf_parser, 'extract_by_regex')
        assert hasattr(docx_parser, 'extract_by_regex')
